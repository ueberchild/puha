# import os
# os.system('pip3 install -r requirements.txt')
# import platform
# if platform.system() == 'Windows': os.system('doc_to_docx.bat')
# else: os.system('bash doc_to_docx.sh')

import pandas as pd, re, openpyxl, io, csv
from docx import Document

class lesson:
    def __init__(self, program_type, time_borders, auditory, personnel_type, specialization):
        self.program_type = program_type
        self.time_borders = time_borders
        self.personnel_type = personnel_type
        self.auditory = auditory
        self.specialization = specialization

def read_docx_tables(filename, tab_id=None, **kwargs):
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)
    
    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise
'''
def read_excel_without_invis(fname): 
    wb = openpyxl.load_workbook(fname)
    ws = wb.worksheets[0]
    hidden_cols = set()
    for col_letter, col_dimensions in ws.column_dimensions.items():
        if col_dimensions.hidden == True:
            hidden_cols.add(col_letter)
    return pd.read_excel(fname).drop(list(hidden_cols), axis = 1)
'''

a_1 = pd.read_excel('data/Приложение №1.xlsx').dropna(how='all').dropna(how='all',axis=1)
a_1.drop(a_1.columns[0:2], axis = 1, inplace = True)

lessons = []

for (month, col_d) in a_1.iteritems():
    if col_d.any() and 'Неделя' in str(col_d.iloc[0]):
        for plan_num, week_plan in enumerate(col_d):
            if type(week_plan) == str and week_plan:
                week_lessons = week_plan.split('\n\n')
                for w_l in week_lessons:
                    w_l_info = w_l.split('\n')
                    if len(w_l_info) == 3 and type(a_1.iloc[plan_num, 0]) == str:
                        time_borders = re.findall('\d*\.\d*', w_l_info[1].replace('..', '.'))
                        personnel_type = 'av' if plan_num < 6 else 'notav'
                        if ' ' in w_l_info[2]:
                            auditory = w_l_info[2].split(' ')[-1]
                        elif '.' in w_l_info[2]:
                            auditory = w_l_info[2].split('.')[-1]
                        lessons.append(lesson(w_l_info[0], time_borders, auditory, personnel_type, a_1.iloc[plan_num, 0]))

curriculum_pars = pd.read_excel('data/Приложение №2.xlsx', sheet_name = 0)
auditory_pars = pd.read_excel('data/Приложение №2.xlsx', sheet_name = 1)
lecturers_pars = pd.read_excel('data/Приложение №2.xlsx', sheet_name = 2)

for c in tuple(curriculum_pars.columns):
    for row in range(len(curriculum_pars)):
        if type(curriculum_pars.at[row, c]) == int:
            curriculum_pars.at[row, c] = max(0, curriculum_pars.at[row, c])

# course = read_docx_tables("data/Приложение №3/Организация наземного обслуживания воздушных судов.docx")
